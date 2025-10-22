import os
from pathlib import Path
from typing import Optional, Union

import pandas as pd
from PyPDF2 import PdfReader, PdfWriter
from docx import Document

from agno.tools import Toolkit
from agno.models.google import Gemini
from agno.utils.log import log_info, log_error


class DocumentProcessingTool(Toolkit):
    def __init__(self, llm: Optional[Gemini] = None):
        super().__init__()
        self.llm = llm or Gemini()

    def read_pdf(self, file_path: Union[str, Path]) -> str:
        try:
            reader = PdfReader(str(file_path))
            text = ""
            for page in reader.pages:
                text += page.extract_text() or ""
            return text
        except Exception as e:
            log_error(f"Error reading PDF: {e}")
            return ""

    def read_docx(self, file_path: Union[str, Path]) -> str:
        try:
            doc = Document(str(file_path))
            return "\n".join([para.text for para in doc.paragraphs])
        except Exception as e:
            log_error(f"Error reading DOCX: {e}")
            return ""

    def read_excel(self, file_path: Union[str, Path], sheet_name: Optional[str] = None) -> pd.DataFrame:
        try:
            df = pd.read_excel(file_path, sheet_name=sheet_name)
            return df
        except Exception as e:
            log_error(f"Error reading Excel file: {e}")
            return pd.DataFrame()

    def read_txt(self, file_path: Union[str, Path]) -> str:
        try:
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        except Exception as e:
            log_error(f"Error reading TXT: {e}")
            return ""

    def write_pdf(self, text: str, output_path: Union[str, Path]):
        try:
            writer = PdfWriter()
            from PyPDF2 import PageObject

            page = PageObject.create_blank_page(width=612, height=792)
            page.merge_text(text, 72, 720)
            writer.add_page(page)
            with open(output_path, "wb") as f:
                writer.write(f)
            log_info(f"PDF written to {output_path}")
        except Exception as e:
            log_error(f"Error writing PDF: {e}")

    def write_docx(self, text: str, output_path: Union[str, Path]):
        try:
            doc = Document()
            for line in text.split("\n"):
                doc.add_paragraph(line)
            doc.save(output_path)
            log_info(f"DOCX written to {output_path}")
        except Exception as e:
            log_error(f"Error writing DOCX: {e}")

    def write_excel(self, df: pd.DataFrame, output_path: Union[str, Path], sheet_name: str = "Sheet1"):
        try:
            df.to_excel(output_path, index=False, sheet_name=sheet_name)
            log_info(f"Excel file written to {output_path}")
        except Exception as e:
            log_error(f"Error writing Excel: {e}")

    def write_txt(self, text: str, output_path: Union[str, Path]):
        try:
            with open(output_path, "w", encoding="utf-8") as f:
                f.write(text)
            log_info(f"TXT written to {output_path}")
        except Exception as e:
            log_error(f"Error writing TXT: {e}")

    
    def process_file(self, file_path: Union[str, Path], output_path: Union[str, Path], sheet_name: Optional[str] = None):
        ext = Path(file_path).suffix.lower()

        if ext == ".pdf":
            content = self.read_pdf(file_path)
        elif ext in [".doc", ".docx"]:
            content = self.read_docx(file_path)
        elif ext in [".xls", ".xlsx"]:
            df = self.read_excel(file_path, sheet_name=sheet_name)
            content = df.to_csv(index=False)
        elif ext == ".txt":
            content = self.read_txt(file_path)
        else:
            log_error(f"Unsupported file type: {ext}")
            return

        if not content:
            log_error(f"No content read from {file_path}")
            return

        # Send content to Gemini LLM
        try:
            modified_content = self.llm.chat(content).content
        except Exception as e:
            log_error(f"LLM processing failed: {e}")
            modified_content = content

        # Write back
        if ext == ".pdf":
            self.write_pdf(modified_content, output_path)
        elif ext in [".doc", ".docx"]:
            self.write_docx(modified_content, output_path)
        elif ext in [".xls", ".xlsx"]:
            try:
                df_modified = pd.read_csv(pd.compat.StringIO(modified_content))
            except Exception:
                df_modified = pd.DataFrame([modified_content.split("\n")])
            self.write_excel(df_modified, output_path, sheet_name=sheet_name or "Sheet1")
        elif ext == ".txt":
            self.write_txt(modified_content, output_path)
